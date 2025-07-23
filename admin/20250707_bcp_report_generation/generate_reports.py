#!/usr/bin/env python3

# The MIT License
# Copyright (c) 2025 Adrian Tan <adrian_tan@nparks.gov.sg>
# Permission is hereby granted, free of charge, to any person obtaining a copy
# of this software and associated documentation files (the 'Software'), to deal
# in the Software without restriction, including without limitation the rights
# to use, copy, modify, merge, publish, distribute, sublicense, and/or sell
# copies of the Software, and to permit persons to whom the Software is
# furnished to do so, subject to the following conditions:
# The above copyright notice and this permission notice shall be included in
# all copies or substantial portions of the Software.
# THE SOFTWARE IS PROVIDED 'AS IS', WITHOUT WARRANTY OF ANY KIND, EXPRESS OR
# IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,
# FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE
# AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER
# LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM
# OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN
# THE SOFTWARE.

import os
import click
import openpyxl
import tkinter as tk
from tkinter import messagebox
import docx
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

@click.command()
@click.option(
    "-w",
    "--working_dir",
    default=os.getcwd(),
    show_default=True,
    help="working directory",
)
@click.option(
    "-d",
    "--data_dir",
    required = True,
    show_default=True,
    help="data directory",
)
def main(working_dir, data_dir):
    """
    Combine data from 3 BCP files and generate various reports.

    e.g. generate_bcp_reports.py
    """
    working_dir = os.path.abspath(working_dir)
    data_dir = os.path.abspath(data_dir)
    reports_dir = f"{working_dir}/reports"
    print("\t{0:<20} :   {1:<10}".format("working dir", working_dir))
    print("\t{0:<20} :   {1:<10}".format("data dir", data_dir))
    print("\t{0:<20} :   {1:<10}".format("reports dir", reports_dir))



    bcp = BCP(data_dir, working_dir)
    bcp.initialise_bcp_files()
    bcp.generate_reports()


    ################################
    # GRAPHICAL USER INTERFACE (GUI)
    ################################

    def say_hello():
        messagebox.showinfo("Hello", "Hello, World!")

    #create main window
    root = tk.Tk()
    root.title("BCP Report Generation")
    root.geometry("400x200")

    label = tk.Label(root, text="BCP Report Generation in Progress...")
    label.pack(pady=20)

    entry = tk.Entry(root, width=50)
    entry.pack(pady=10)

    button = tk.Button(root, text="Generate Reports", command=say_hello)
    button.pack(pady=10)


    #root.mainloop()

class BCP(object):
    def __init__(self, bcp_dir, working_dir):
        self.bcp_dir = bcp_dir
        self.working_dir = working_dir
        self.reports_dir = f"{self.working_dir}/reports"

        # create directories
        try:
            os.makedirs(self.reports_dir, exist_ok=True)
        except OSError as error:
            print(f"{error.filename} cannot be created")


    def initialise_bcp_files(self):
        print(f"initializing bcp files")

        #read combined lab files
        # CAVS_Lab_combined.xlsx
        # bcp_files/CAVS\ FormSG\ sample\ submission\ spreadsheets
        print(f"==========================")
        print(f"reading combined lab files")
        print(f"==========================")
        workbook = openpyxl.load_workbook(
            filename=f"{self.bcp_dir}/CAVS_Lab_combined.xlsx",
            data_only=True,
        )
        print(workbook.sheetnames)

        print(f"=======================")
        print(f"reading lab tests files")
        print(f"=======================")
        lab_tests_workbook = openpyxl.load_workbook(
            filename=f"{self.bcp_dir}/BCP_CAVS Laboratory Tests_24 Feb 2025.xlsx",
            data_only=True,
        )
        print(lab_tests_workbook.sheetnames)

        print(f"=======================================")
        print(f"iterate through lab results directories")
        print(f"=======================================")
        lab_results_dir = f"{self.bcp_dir}/CAVS FormSG sample submission spreadsheets"
        for file_name in os.listdir(lab_results_dir):
            print(file_name)

    def set_table_borders(self, table):
        """
        Apply borders to all cells in a python-docx table.
        """
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')

                for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '4')  # 4 eighths of a point = 0.5pt
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), '000000')
                    tcBorders.append(border)

                tcPr.append(tcBorders)

    def generate_reports(self):
        print(f"==========================")
        print(f"Generating reports for BCP")
        print(f"==========================")
        # Create a new document
        doc = docx.Document()

        # Add a title
        doc.add_heading('My Report', level=0)

        # Add a heading
        doc.add_heading('Introduction', level=1)

        # Add a paragraph
        doc.add_paragraph('This document is generated using the python-docx library.')

        # Add a sub-heading
        doc.add_heading('Data Summary', level=2)

        # Add another paragraph
        doc.add_paragraph('Below is a table summarizing the data:')

        # Add a table with 3 rows and 3 columns
        table = doc.add_table(rows=3, cols=3)

        # Add table headers
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'ID'
        hdr_cells[1].text = 'Name'
        hdr_cells[2].text = 'Score'

        # Populate data rows
        table.rows[1].cells[0].text = '1'
        table.rows[1].cells[1].text = 'Alice'
        table.rows[1].cells[2].text = '85'

        table.rows[2].cells[0].text = '2'
        table.rows[2].cells[1].text = 'Bob'
        table.rows[2].cells[2].text = '90'

        self.set_table_borders(table)

        # Save the document
        print(f"Saving report to {self.reports_dir}/example_report.docx")
        doc.save(f'{self.reports_dir}/example_report.docx')        



if __name__ == "__main__":
    main() # type: ignore
